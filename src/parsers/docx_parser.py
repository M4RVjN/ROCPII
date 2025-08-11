# src/parsers/docx_parser.py
import logging
import pathlib
import docx
from src.shared_data_model import FileContext, FileStatus
from src.parsers.base_parser import BaseParser

class DocxParser(BaseParser):
    def supports(self, mime_type: str) -> bool:
        return mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    def parse(self, file_path: pathlib.Path) -> tuple[FileContext, str]:
        ctx_args = {"file_path": file_path, "mime_type": self.supports.__annotations__['mime_type'], "file_size_bytes": file_path.stat().st_size}
        try:
            document = docx.Document(file_path)
            all_text = []
            def extract_from(container):
                if container is None: return
                for para in container.paragraphs: all_text.append(para.text)
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells: all_text.append(cell.text)
            extract_from(document)
            for section in document.sections:
                extract_from(section.header)
                extract_from(section.footer)
            full_text = "\n".join(filter(None, all_text))
            ctx = FileContext(**ctx_args, status=FileStatus.COMPLETED)
            return ctx, full_text
        except Exception as e:
            msg = f"解析 DOCX 檔案時發生錯誤: {e}"
            if 'encrypted' in str(e).lower(): msg = "檔案可能已加密或已損毀，無法解析。"
            logging.error(f"'{file_path.name}': {msg}", exc_info=True)
            ctx = FileContext(**ctx_args, status=FileStatus.ERROR, error_message=msg)
            return ctx, ""